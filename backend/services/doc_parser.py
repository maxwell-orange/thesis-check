"""
Word document parser using python-docx
Extracts structured information from .docx files
"""
from dataclasses import dataclass
from typing import List, Optional, Dict, Any
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re


@dataclass
class ParagraphInfo:
    """Paragraph information"""
    text: str
    style_name: str
    font_name: str
    font_size: Pt
    font_size_pt: float
    bold: bool
    italic: bool
    alignment: str
    first_line_indent: float  # in cm
    line_spacing: float
    space_before: Pt
    space_after: Pt
    paragraph_index: int

    def to_dict(self) -> Dict:
        return {
            'text': self.text,
            'style_name': self.style_name,
            'font_name': self.font_name,
            'font_size_pt': self.font_size_pt,
            'bold': self.bold,
            'italic': self.italic,
            'alignment': self.alignment,
            'first_line_indent': self.first_line_indent,
            'line_spacing': self.line_spacing,
            'space_before': self.space_before,
            'space_after': self.space_after,
            'paragraph_index': self.paragraph_index
        }


@dataclass
class TableInfo:
    """Table information"""
    table_index: int
    row_count: int
    col_count: int
    caption: Optional[str]
    position: str  # 'above' or 'below' for caption position


@dataclass
class PageSettings:
    """Page settings"""
    paper_size: str
    orientation: str
    margin_top: float  # cm
    margin_bottom: float  # cm
    margin_left: float  # cm
    margin_right: float  # cm

    def to_dict(self) -> Dict:
        return {
            'paper_size': self.paper_size,
            'orientation': self.orientation,
            'margin_top': self.margin_top,
            'margin_bottom': self.margin_bottom,
            'margin_left': self.margin_left,
            'margin_right': self.margin_right
        }


@dataclass
class DocumentStructure:
    """Document structure"""
    title: Optional[str]
    abstract_cn: Optional[str]
    abstract_en: Optional[str]
    keywords_cn: Optional[List[str]]
    keywords_en: Optional[List[str]]
    chapters: List[Dict[str, Any]]
    references: List[str]
    has_acknowledgement: bool
    has_appendix: bool


class DocParser:
    """Word document parser"""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
        self.paragraphs: List[ParagraphInfo] = []
        self.tables: List[TableInfo] = []
        self.page_settings: Optional[PageSettings] = None
        self.structure: Optional[DocumentStructure] = None

    def parse(self) -> 'DocParser':
        """Parse the entire document"""
        self.extract_paragraphs()
        self.extract_tables()
        self.extract_page_settings()
        self.extract_structure()
        return self

    def extract_paragraphs(self) -> List[ParagraphInfo]:
        """Extract all paragraphs with formatting info"""
        self.paragraphs = []
        list_idx = 0  # Index in self.paragraphs list

        for idx, para in enumerate(self.doc.paragraphs):
            # Skip empty paragraphs
            if not para.text.strip():
                continue

            # Get run formatting (from first run if exists)
            run = para.runs[0] if para.runs else None

            # Get font properties
            font_name = "Unknown"
            font_size = Pt(12)
            font_size_pt = 12.0
            bold = False
            italic = False

            if run and run.font:
                # Try to get font name
                if run.font.name:
                    font_name = run.font.name
                else:
                    # Try to get from East Asian font
                    try:
                        r = run._element
                        rPr = r.find(qn('w:rPr'))
                        if rPr is not None:
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is not None:
                                font_name = rFonts.get(qn('w:eastAsia')) or rFonts.get(qn('w:ascii')) or "Unknown"
                    except:
                        pass

                # Get font size
                if run.font.size:
                    font_size = run.font.size
                    font_size_pt = font_size.pt if hasattr(font_size, 'pt') else 12.0

                bold = bool(run.font.bold)
                italic = bool(run.font.italic)

            # Get paragraph properties
            alignment = "left"
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                alignment = "center"
            elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                alignment = "right"
            elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                alignment = "justify"

            # Get indentation
            first_line_indent = 0.0
            if para.paragraph_format.first_line_indent:
                try:
                    first_line_indent = para.paragraph_format.first_line_indent.cm
                except:
                    pass

            # Get line spacing
            line_spacing = 1.0
            if para.paragraph_format.line_spacing:
                line_spacing = para.paragraph_format.line_spacing

            # Get spacing before/after
            space_before = Pt(0)
            space_after = Pt(0)
            if para.paragraph_format.space_before:
                try:
                    space_before = para.paragraph_format.space_before.pt
                except:
                    pass
            if para.paragraph_format.space_after:
                try:
                    space_after = para.paragraph_format.space_after.pt
                except:
                    pass

            para_info = ParagraphInfo(
                text=para.text.strip(),
                style_name=para.style.name if para.style else "Normal",
                font_name=font_name,
                font_size=font_size,
                font_size_pt=font_size_pt,
                bold=bold,
                italic=italic,
                alignment=alignment,
                first_line_indent=first_line_indent,
                line_spacing=line_spacing,
                space_before=space_before,
                space_after=space_after,
                paragraph_index=list_idx
            )

            self.paragraphs.append(para_info)
            list_idx += 1

        return self.paragraphs

    def extract_tables(self) -> List[TableInfo]:
        """Extract table information"""
        self.tables = []

        for idx, table in enumerate(self.doc.tables):
            # Try to find table caption
            caption = None
            position = "above"  # Default for tables

            # Check text before and after table
            table_element = table._element
            body = self.doc.element.body

            table_info = TableInfo(
                table_index=idx,
                row_count=len(table.rows),
                col_count=len(table.columns),
                caption=caption,
                position=position
            )
            self.tables.append(table_info)

        return self.tables

    def extract_page_settings(self) -> PageSettings:
        """Extract page settings"""
        section = self.doc.sections[0] if self.doc.sections else None

        if section:
            # Try to get page dimensions
            try:
                page_width = section.page_width.cm if section.page_width else 21.0
                page_height = section.page_height.cm if section.page_height else 29.7

                # A4 is 21cm x 29.7cm
                paper_size = "A4" if abs(page_width - 21.0) < 0.5 and abs(page_height - 29.7) < 0.5 else "Unknown"

                orientation = "portrait" if page_height > page_width else "landscape"
            except:
                paper_size = "Unknown"
                orientation = "portrait"

            # Get margins
            try:
                margin_top = section.top_margin.cm if section.top_margin else 2.5
                margin_bottom = section.bottom_margin.cm if section.bottom_margin else 2.0
                margin_left = section.left_margin.cm if section.left_margin else 2.5
                margin_right = section.right_margin.cm if section.right_margin else 2.0
            except:
                margin_top = 2.5
                margin_bottom = 2.0
                margin_left = 2.5
                margin_right = 2.0
        else:
            paper_size = "A4"
            orientation = "portrait"
            margin_top = 2.5
            margin_bottom = 2.0
            margin_left = 2.5
            margin_right = 2.0

        self.page_settings = PageSettings(
            paper_size=paper_size,
            orientation=orientation,
            margin_top=margin_top,
            margin_bottom=margin_bottom,
            margin_left=margin_left,
            margin_right=margin_right
        )

        return self.page_settings

    def extract_structure(self) -> DocumentStructure:
        """Extract document structure (chapters, sections, etc.)"""
        chapters = []
        current_chapter = None
        references = []
        abstract_cn = None
        abstract_en = None
        keywords_cn = None
        keywords_en = None
        has_acknowledgement = False
        has_appendix = False

        # State machine for parsing
        state = "start"

        for para in self.paragraphs:
            text = para.text

            # Detect chapters (e.g., "第一章", "第1章", "1. 绪论", "一、", "二、")
            chapter_match = re.match(r'^第[一二三四五六七八九十\d]+章', text)
            chinese_chapter_match = re.match(r'^[一二三四五六七八九十]+[、\.\s]', text)
            section_match = re.match(r'^\d+\.\d+\s+', text)
            subsection_match = re.match(r'^\d+\.\d+\.\d+\s+', text)

            if chapter_match or chinese_chapter_match:
                current_chapter = {
                    'title': text,
                    'paragraph_index': para.paragraph_index,
                    'sections': [],
                    'font_size': para.font_size_pt,
                    'font_name': para.font_name,
                    'bold': para.bold,
                    'alignment': para.alignment
                }
                chapters.append(current_chapter)
                state = "chapter"
            elif section_match and current_chapter:
                section = {
                    'title': text,
                    'paragraph_index': para.paragraph_index,
                    'subsections': [],
                    'font_size': para.font_size_pt,
                    'font_name': para.font_name,
                    'bold': para.bold
                }
                current_chapter['sections'].append(section)
                state = "section"
            elif subsection_match and current_chapter and current_chapter['sections']:
                subsection = {
                    'title': text,
                    'paragraph_index': para.paragraph_index,
                    'font_size': para.font_size_pt,
                    'font_name': para.font_name,
                    'bold': para.bold,
                    'first_line_indent': para.first_line_indent
                }
                current_chapter['sections'][-1]['subsections'].append(subsection)
                state = "subsection"

            # Detect abstract
            if "摘要" in text and len(text) < 10 and abstract_cn is None:
                # Next paragraph should be abstract content
                idx = self.paragraphs.index(para) if para in self.paragraphs else -1
                if idx >= 0 and idx + 1 < len(self.paragraphs):
                    abstract_cn = self.paragraphs[idx + 1].text

            # Detect English abstract
            if "abstract" in text.lower() and len(text) < 20 and abstract_en is None:
                idx = self.paragraphs.index(para) if para in self.paragraphs else -1
                if idx >= 0 and idx + 1 < len(self.paragraphs):
                    abstract_en = self.paragraphs[idx + 1].text

            # Detect references
            if re.search(r'参考.?文献', text) and len(text) < 10:
                state = "references"
            elif state == "references" and text.strip():
                if re.match(r'^\[\d+\]', text):
                    references.append(text)

            # Detect acknowledgement
            if re.search(r'致谢', text) and len(text) < 10:
                has_acknowledgement = True

            # Detect appendix
            if re.search(r'附录', text) and len(text) < 10:
                has_appendix = True

        # Filter out TOC entries - chapters with very few paragraphs (likely directory entries)
        # Keep only chapters that have substantial content (> 5 paragraphs)
        filtered_chapters = []
        for i, ch in enumerate(chapters):
            start_idx = ch['paragraph_index']
            if i + 1 < len(chapters):
                end_idx = chapters[i + 1]['paragraph_index']
            else:
                end_idx = len(self.paragraphs)
            para_count = end_idx - start_idx

            # Keep chapter if it has more than 5 paragraphs or it's the only chapter
            if para_count > 5 or len(chapters) == 1:
                filtered_chapters.append(ch)

        chapters = filtered_chapters

        # Try to get title from first non-empty paragraph
        title = None
        for para in self.paragraphs[:5]:  # Check first 5 paragraphs
            if para.alignment == "center" and len(para.text) > 10:
                title = para.text
                break

        self.structure = DocumentStructure(
            title=title,
            abstract_cn=abstract_cn,
            abstract_en=abstract_en,
            keywords_cn=keywords_cn,
            keywords_en=keywords_en,
            chapters=chapters,
            references=references,
            has_acknowledgement=has_acknowledgement,
            has_appendix=has_appendix
        )

        return self.structure

    def get_full_text(self) -> str:
        """Get full document text"""
        return "\n".join([p.text for p in self.doc.paragraphs])

    def get_text_by_style(self, style_name: str) -> List[str]:
        """Get text by style name"""
        return [p.text for p in self.doc.paragraphs
                if p.style and p.style.name == style_name]

    def get_chapter_content(self, chapter_index: int) -> str:
        """Get content of a specific chapter"""
        if not self.structure or chapter_index >= len(self.structure.chapters):
            return ""

        chapter = self.structure.chapters[chapter_index]
        start_idx = chapter.get('paragraph_index', 0)

        if chapter_index + 1 < len(self.structure.chapters):
            end_idx = self.structure.chapters[chapter_index + 1].get('paragraph_index', len(self.paragraphs))
        else:
            end_idx = len(self.paragraphs)

        chapter_paragraphs = [p.text for p in self.paragraphs[start_idx:end_idx]]
        return "\n".join(chapter_paragraphs)

    def extract_citations(self) -> List[Dict[str, Any]]:
        """
        Extract citation markers from body text like [1], [2-5], [3,6,8]
        Returns list of dicts with citation info
        """
        citations = []
        citation_pattern = re.compile(r'\[(\d+(?:\s*[-,\s]\s*\d+)*)\]')

        for para in self.paragraphs:
            text = para.text
            matches = citation_pattern.findall(text)

            for match in matches:
                # Parse citation like "1", "2-5", "3,6,8"
                ref_numbers = []
                parts = re.split(r'\s*,\s*|\s+', match)

                for part in parts:
                    part = part.strip()
                    if '-' in part or '–' in part:
                        # Range like "2-5"
                        try:
                            start, end = re.split(r'[-–]', part)
                            ref_numbers.extend(range(int(start), int(end) + 1))
                        except (ValueError, TypeError):
                            continue
                    else:
                        try:
                            ref_numbers.append(int(part))
                        except ValueError:
                            continue

                for ref_num in ref_numbers:
                    citations.append({
                        'reference_number': ref_num,
                        'context': text.strip()[:100] + '...' if len(text) > 100 else text.strip(),
                        'full_match': f'[{match}]'
                    })

        return citations

    def get_references_section(self) -> str:
        """Get the references section text"""
        if not self.structure or not self.structure.references:
            return ""
        return "\n".join(self.structure.references)

    def get_references_with_indices(self) -> List[Dict[str, str]]:
        """
        Get references with their indices parsed
        Returns list of dicts with index and text
        """
        refs = []
        for ref_text in self.structure.references if self.structure else []:
            # Try to extract index from [1], [2], etc.
            match = re.match(r'^\[(\d+)\]\s*(.*)$', ref_text.strip())
            if match:
                refs.append({
                    'index': int(match.group(1)),
                    'text': match.group(2),
                    'full_text': ref_text
                })
            else:
                refs.append({
                    'index': None,
                    'text': ref_text,
                    'full_text': ref_text
                })
        return refs
